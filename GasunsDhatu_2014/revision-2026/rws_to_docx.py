# -*- coding: utf-8 -*-
"""H385 — build the review .docx from the source .mdx + rws_edits.jsonl.

Renders the whole GasunsDhatu 2026 book as a Word document in which every
paragraph an RWS pass rewrote is:
  * shown in its NEW (edited) form,
  * filled with a yellow highlight, and
  * annotated with a Word comment  «Было: <original text>»  plus the finding
    ids and the reviewer style(s) that motivated the edit,
so the author can review every change in Google Docs and roll any of them back
by hand (the "Было" text is the exact pre-edit paragraph).

Usage:
    python rws_to_docx.py            # -> revision-2026/GasunsDhatu_2026_RWS_review.docx

Requires python-docx >= 1.1 (Document.add_comment). Stdlib + python-docx only.
"""
import json, sys, os
sys.stdout.reconfigure(encoding='utf-8')
sys.stderr.reconfigure(encoding='utf-8')
from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx.shared import Pt

BASE = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))  # GasunsDhatu_2014
REV = os.path.join(BASE, 'revision-2026')
EDITS = os.path.join(REV, 'rws_edits.jsonl')
OUT = os.path.join(REV, 'GasunsDhatu_2026_RWS_review.docx')

# render order: main text, then the three appendix articles
FILE_ORDER = [
    '02_gasuns-dhatu-PhD-text2.mdx',
    'Морфонологическая запись глагольных корней санскрита.mdx',
    'О записи омонимии корней в словарях древнеиндийского языка.mdx',
    'Распределение рядов согласных.mdx',
]


def load_edits():
    edits = {}  # (file, before_text) -> record
    if not os.path.exists(EDITS):
        return edits
    with open(EDITS, encoding='utf-8') as f:
        for ln in f:
            ln = ln.strip()
            if not ln:
                continue
            r = json.loads(ln)
            # key on `after` — this runs against the ALREADY-EDITED .mdx, whose
            # paragraphs now hold the new text; we highlight those and put the
            # original in the comment.
            edits[(r['file'], r['after'])] = r
    return edits


def split_blocks(text):
    """Yield ('blank'|'block', text) preserving structure by blank-line splits."""
    lines = text.split('\n')
    buf, i, n = [], 0, len(lines)
    blocks = []
    for line in lines:
        if line.strip() == '':
            if buf:
                blocks.append('\n'.join(buf))
                buf = []
            blocks.append(None)  # blank marker
        else:
            buf.append(line)
    if buf:
        blocks.append('\n'.join(buf))
    return blocks


def strip_frontmatter(text):
    if text.startswith('---'):
        end = text.find('\n---', 3)
        if end != -1:
            nl = text.find('\n', end + 1)
            text = text[nl + 1:] if nl != -1 else ''
    return text


def add_block(doc, block, edit):
    """Render one non-blank block; if edit is not None, highlight + comment."""
    first = block.split('\n', 1)[0]
    # headings
    if first.lstrip().startswith('#') and edit is None:
        level = len(first) - len(first.lstrip('#').lstrip()) if False else first.count('#', 0, len(first) - len(first.lstrip('#')))
        hlevel = min(max(first[:8].count('#'), 1), 4)
        doc.add_heading(block.lstrip('#').strip(), level=hlevel)
        return
    p = doc.add_paragraph()
    runs = []
    parts = block.split('\n')
    for k, part in enumerate(parts):
        r = p.add_run(part)
        if edit is not None:
            r.font.highlight_color = WD_COLOR_INDEX.YELLOW
        runs.append(r)
        if k != len(parts) - 1:
            br = p.add_run()
            br.add_break()
    if edit is not None:
        fids = ', '.join('#' + str(i) for i in edit.get('finding_ids', []))
        styles = ', '.join(edit.get('styles', []))
        comment = 'Было:\n' + edit['before']
        tail = []
        if fids:
            tail.append('находки: ' + fids)
        if styles:
            tail.append('стиль: ' + styles)
        if tail:
            comment += '\n\n[RWS — ' + '; '.join(tail) + ']'
        try:
            doc.add_comment(runs, text=comment, author='RWS-2026', initials='RWS')
        except Exception as e:
            # fall back to an inline bracket note so nothing is lost
            note = p.add_run('  [Было: ' + edit['before'][:200] + ']')
            note.italic = True


def main():
    edits = load_edits()
    doc = Document()
    doc.add_heading('GasunsDhatu 2026 — просмотр RWS-правок', level=0)
    intro = doc.add_paragraph()
    intro.add_run(
        'Жёлтой заливкой отмечен каждый переработанный абзац; в комментарии к нему '
        '(«Было: …») — исходный текст до правки, номера RWS-находок и стиль-рецензент. '
        'Чтобы отклонить правку, верните текст из комментария вручную. '
        'Данные и цифры в этой задаче не менялись.').italic = True
    applied = 0
    for fn in FILE_ORDER:
        path = os.path.join(BASE, fn)
        if not os.path.exists(path):
            continue
        doc.add_heading('Файл: ' + fn, level=1)
        text = strip_frontmatter(open(path, encoding='utf-8').read())
        for block in split_blocks(text):
            if block is None:
                continue
            edit = edits.get((fn, block))
            if edit is not None:
                applied += 1
            add_block(doc, block, edit)
    doc.save(OUT)
    print('edits in jsonl:', len(edits))
    print('edited paragraphs rendered:', applied)
    print('saved:', OUT)
    if applied != len(edits):
        print('WARNING: %d jsonl edits did not match a source paragraph verbatim'
              % (len(edits) - applied))


if __name__ == '__main__':
    main()
